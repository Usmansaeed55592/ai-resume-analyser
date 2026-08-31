"""
Converts the AI-regenerated resume text into downloadable files (DOCX / PDF).
"""

import io
import re
import unicodedata
import docx
from fpdf import FPDF


def _strip_markdown_artifacts(text: str) -> str:
    """Clean up markdown syntax the AI might still emit (headers, bold,
    tables, horizontal rules) so the PDF/DOCX renderer never shows raw
    '**', '##', '|' characters - only clean plain text."""
    cleaned_lines = []
    for raw_line in text.splitlines():
        stripped = raw_line.strip()

        # Drop markdown horizontal rules / stray separator lines ("---", "***")
        if re.fullmatch(r"[-*_]{2,}", stripped) or stripped == "-":
            continue

        # Drop markdown table separator rows, e.g. |---|---| or | :--- | ---: |
        if "|" in stripped and re.fullmatch(r"\|?[\s:|-]+\|?", stripped):
            continue

        # Headers: "## Something" -> "SOMETHING" (so heading detection catches it)
        header_match = re.match(r"^#{1,6}\s*(.+)$", stripped)
        if header_match:
            cleaned_lines.append(header_match.group(1).strip(" *_").upper())
            continue

        line = raw_line.rstrip()

        # Markdown table rows: "| A | B | C |" -> "A | B | C" (drop the pipes' padding)
        if stripped.startswith("|") and stripped.endswith("|") and stripped.count("|") >= 2:
            cells = [c.strip(" *_") for c in stripped.strip("|").split("|") if c.strip()]
            line = "  |  ".join(cells) if cells else ""
            if not line:
                continue

        # Remove bold/italic markers: **text**, __text__, *text* -> text
        line = re.sub(r"\*\*(.+?)\*\*", r"\1", line)
        line = re.sub(r"__(.+?)__", r"\1", line)
        line = re.sub(r"(?<!\w)\*(?!\s)(.+?)(?<!\s)\*(?!\w)", r"\1", line)

        # Strip any stray leftover leading '#' marks
        line = re.sub(r"^\s*#{1,6}\s*", "", line)

        cleaned_lines.append(line)

    return "\n".join(cleaned_lines)


def _sanitize_text(text: str) -> str:
    """Replace characters that core PDF fonts (Helvetica) can't render.

    Previously this only mapped a couple of dash/quote variants and fell
    back to latin-1 "replace" for anything else - which silently turns ANY
    unmapped character (other dash variants, bullets, primes, etc.) into a
    literal "?" in the PDF. That's what produced "server?rendered",
    "role?based" etc. This version covers every common typographic
    character an LLM tends to emit, then degrades gracefully (accented
    letters -> base letter, anything else truly unrenderable -> "-")
    instead of ever printing "?".
    """
    replacements = {
        # bullets
        "\u2022": "-", "\u2023": "-", "\u25e6": "-", "\u2043": "-", "\u00b7": "-",
        # dashes/hyphens - this is the category that was breaking, LLMs use many variants
        "\u2010": "-", "\u2011": "-", "\u2012": "-", "\u2013": "-",
        "\u2014": "-", "\u2015": "-", "\u2212": "-",
        "\ufe58": "-", "\ufe63": "-", "\uff0d": "-",
        # quotes/apostrophes
        "\u2018": "'", "\u2019": "'", "\u201a": "'", "\u201b": "'",
        "\u201c": '"', "\u201d": '"', "\u201e": '"', "\u201f": '"',
        "\u2032": "'", "\u2033": '"',
        # misc
        "\u2026": "...",
        "\u00a0": " ", "\u2000": " ", "\u2001": " ", "\u2002": " ",
        "\u2003": " ", "\u2004": " ", "\u2005": " ", "\u2006": " ",
        "\u2007": " ", "\u2008": " ", "\u2009": " ", "\u200a": " ",
        "\u200b": "", "\ufeff": "",
    }
    for bad, good in replacements.items():
        text = text.replace(bad, good)

    # Decompose accented letters (e.g. "e with accent" -> "e" + mark) so
    # words with accents degrade to their plain-letter form instead of
    # vanishing or turning into "?".
    text = unicodedata.normalize("NFKD", text)
    text = "".join(ch for ch in text if not unicodedata.combining(ch))

    # Anything still outside Latin-1 (core PDF fonts can't render it) becomes
    # a plain hyphen instead of "?" so the resume never shows garbled glyphs.
    out_chars = []
    for ch in text:
        try:
            ch.encode("latin-1")
            out_chars.append(ch)
        except UnicodeEncodeError:
            out_chars.append("-")
    return "".join(out_chars)

def _break_word_to_fit(pdf, word: str, max_width: float) -> str:
    """Break a single word into pieces that each actually fit max_width,
    measured with the PDF's current font (exact width, not a char guess)."""
    if pdf.get_string_width(word) <= max_width:
        return word

    pieces = []
    current = ""
    for ch in word:
        candidate = current + ch
        if current and pdf.get_string_width(candidate) > max_width:
            pieces.append(current)
            current = ch
        else:
            current = candidate
    if current:
        pieces.append(current)
    return " ".join(pieces)


def _wrap_long_words(pdf, text: str, max_width: float) -> str:
    """Ensure no single 'word' (unbroken run of non-space characters) is
    wider than max_width, using the PDF's active font to measure it."""
    return " ".join(
        _break_word_to_fit(pdf, word, max_width) if word else word
        for word in text.split(" ")
    )


def _safe_multi_cell(pdf, h: float, text: str, align: str = "L") -> None:
    """multi_cell with a fallback: if it still can't fit (edge case), shrink
    the font size step by step and retry rather than crashing the export."""
    original_size = pdf.font_size_pt
    for size in (original_size, max(original_size - 2, 6), 6):
        pdf.set_font_size(size)
        max_width = pdf.w - pdf.l_margin - pdf.r_margin - 1
        wrapped = _wrap_long_words(pdf, text, max_width)
        try:
            pdf.multi_cell(0, h, wrapped, align=align)
            # fpdf2's single-line fast path can leave the cursor sitting at
            # the right edge instead of resetting to the left margin - force
            # it back so the NEXT multi_cell call always has full width.
            pdf.set_x(pdf.l_margin)
            pdf.set_font_size(original_size)
            return
        except Exception:
            pdf.set_x(pdf.l_margin)
            continue
    # Last resort: skip this line rather than crashing the whole export.
    pdf.set_x(pdf.l_margin)
    pdf.set_font_size(original_size)


def edit_docx_paragraphs(original_file_bytes: bytes, rewritten_paragraphs: list) -> bytes:
    """
    Write AI-rewritten text back into the ORIGINAL uploaded DOCX, preserving
    that file's template (fonts, sizes, colors, bullets, spacing, layout).

    original_file_bytes: bytes of the DOCX exactly as the user uploaded it
    rewritten_paragraphs: list of {"index": int, "text": str} whose "index"
        values line up with paragraph positions in the original document
        (as produced by parser.extract_docx_paragraphs)
    """
    document = docx.Document(io.BytesIO(original_file_bytes))
    replacements = {item["index"]: item["text"] for item in rewritten_paragraphs}

    for idx, paragraph in enumerate(document.paragraphs):
        if idx not in replacements:
            continue
        new_text = replacements[idx]

        if paragraph.runs:
            # Keep the first run's formatting (font/size/bold/color) and put
            # all the new text there; blank out any other runs so text
            # isn't duplicated.
            paragraph.runs[0].text = new_text
            for run in paragraph.runs[1:]:
                run.text = ""
        else:
            paragraph.add_run(new_text)

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def text_to_docx_bytes(resume_text: str, recommendations: list = None) -> bytes:
    """Convert plain resume text into a simple formatted DOCX, returned as bytes.

    recommendations: optional list of genuinely-missing JD keywords/skills
    (e.g. analysis["semantic_missing_keywords"]). These are never inserted
    into the resume body itself (that would be fabrication) - instead they
    are appended as a clearly separate "Career Recommendations" note so the
    user knows exactly what to learn/add before applying.
    """
    resume_text = _strip_markdown_artifacts(resume_text)
    document = docx.Document()

    for line in resume_text.splitlines():
        stripped = line.strip()
        if not stripped:
            continue

        # Treat short all-caps or title-like lines as section headers
        if stripped.isupper() and len(stripped.split()) <= 5:
            document.add_heading(stripped.title(), level=2)
        else:
            document.add_paragraph(stripped)

    if recommendations:
        document.add_page_break()
        heading = document.add_heading("Career Recommendations", level=2)
        document.add_paragraph(
            "This section is NOT part of your resume - it will not be seen by "
            "recruiters or ATS systems. It's a private note for you."
        )
        document.add_paragraph(
            "Your resume was not changed to add these because they weren't found "
            "in it. Consider genuinely learning or gaining experience in the "
            "following before applying to this role:"
        )
        for item in recommendations:
            document.add_paragraph(item, style="List Bullet")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _is_bullet_line(line: str) -> bool:
    return line.startswith(("-", "•", "*", "·"))


def _ensure_space(pdf, needed_mm: float) -> None:
    """If less than needed_mm of vertical space remains before the page's
    bottom margin, start a new page now. Used to keep a heading (or a
    project sub-header) together with at least its first line of content,
    instead of stranding the heading alone at the bottom of a page while
    everything under it flows onto the next page."""
    remaining = pdf.page_break_trigger - pdf.get_y()
    if remaining < needed_mm:
        pdf.add_page()


def text_to_pdf_bytes(resume_text: str, recommendations: list = None) -> bytes:
    """Convert plain resume text into a styled PDF: centered name/title/contact
    header, underlined section headings, bold project sub-headers, and bullets.

    recommendations: optional list of genuinely-missing JD keywords/skills
    (e.g. analysis["semantic_missing_keywords"]). These are never inserted
    into the resume body itself (that would be fabrication) - instead they
    are appended on a separate final page as a "Career Recommendations" note
    so the user knows exactly what to learn/add before applying.
    """
    lines = [l.rstrip() for l in _strip_markdown_artifacts(resume_text).splitlines()]
    while lines and not lines[0].strip():
        lines.pop(0)
    while lines and not lines[-1].strip():
        lines.pop()
    n = len(lines)

    def next_nonempty(i):
        j = i
        while j < n and not lines[j].strip():
            j += 1
        return j

    pdf = FPDF(format="A4")
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=17)
    pdf.set_margins(20, 17, 20)
    pdf.set_font("Helvetica", size=11)

    i = next_nonempty(0)

    # ---- Name ----
    if i < n:
        name = _sanitize_text(lines[i].strip())
        pdf.set_font("Helvetica", "B", 18)
        _safe_multi_cell(pdf, 9, name, align="C")
        i = next_nonempty(i + 1)

    # ---- Job title (short line, no @ / no bullet) ----
    if i < n and lines[i].strip() and "@" not in lines[i] and not _is_bullet_line(lines[i].strip()) \
            and len(lines[i].split()) <= 6 and not lines[i].strip().isupper():
        title = _sanitize_text(lines[i].strip())
        pdf.set_font("Helvetica", "", 12.5)
        pdf.set_text_color(70, 70, 70)
        _safe_multi_cell(pdf, 6.5, title, align="C")
        pdf.set_text_color(0, 0, 0)
        i = next_nonempty(i + 1)

    # ---- Contact line (has @, |, linkedin, github, or a phone-ish pattern) ----
    if i < n:
        lower = lines[i].lower()
        if "@" in lines[i] or "|" in lines[i] or "linkedin" in lower or "github" in lower:
            contact = _sanitize_text(lines[i].strip())
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(60, 60, 60)
            _safe_multi_cell(pdf, 5.5, contact, align="C")
            pdf.set_text_color(0, 0, 0)
            i = next_nonempty(i + 1)

    pdf.ln(4)

    # ---- Body: sections / sub-headers / bullets / paragraphs ----
    while i < n:
        line = lines[i].strip()
        if not line:
            pdf.ln(2.5)
            i += 1
            continue

        stripped = _sanitize_text(line)
        is_heading = (
            stripped.isupper()
            and len(stripped.split()) <= 6
            and not any(ch.isdigit() for ch in stripped)
        )
        is_bullet = _is_bullet_line(stripped)

        j = next_nonempty(i + 1)
        next_is_bullet = j < n and _is_bullet_line(lines[j].strip())

        if is_heading:
            # Keep the heading together with at least its first line of
            # content - never let it sit alone at the bottom of a page.
            _ensure_space(pdf, 24)
            pdf.ln(4.5)
            pdf.set_font("Helvetica", "B", 13)
            _safe_multi_cell(pdf, 6.5, stripped)
            y = pdf.get_y()
            pdf.set_draw_color(40, 40, 40)
            pdf.line(pdf.l_margin, y, pdf.w - pdf.r_margin, y)
            pdf.ln(2.5)
        elif is_bullet:
            text = stripped.lstrip("-•*· ").strip()
            pdf.set_font("Helvetica", "", 11)
            pdf.set_x(pdf.l_margin + 4)
            _safe_multi_cell(pdf, 6, f"-  {text}")
        elif next_is_bullet:
            # A non-bullet line immediately followed by bullets acts like a
            # project/job sub-header - keep it with at least its first
            # bullet, and make it stand out in bold.
            _ensure_space(pdf, 16)
            pdf.set_font("Helvetica", "B", 11)
            _safe_multi_cell(pdf, 6, stripped)
        else:
            pdf.set_font("Helvetica", "", 11)
            _safe_multi_cell(pdf, 6, stripped)

        i += 1

    if recommendations:
        pdf.add_page()
        pdf.set_font("Helvetica", "B", 14)
        _safe_multi_cell(pdf, 8, "Career Recommendations")
        pdf.set_font("Helvetica", "I", 9)
        pdf.set_text_color(100, 100, 100)
        _safe_multi_cell(
            pdf, 5,
            "This page is NOT part of your resume - it will not be seen by "
            "recruiters or ATS systems. It's a private note for you.",
        )
        pdf.set_text_color(0, 0, 0)
        pdf.ln(3)
        pdf.set_font("Helvetica", "", 10.5)
        _safe_multi_cell(
            pdf, 5.5,
            "Your resume was not changed to add these because they weren't found "
            "in it. Consider genuinely learning or gaining experience in the "
            "following before applying to this role:",
        )
        pdf.ln(2)
        pdf.set_font("Helvetica", "", 10.5)
        for item in recommendations:
            pdf.set_x(pdf.l_margin + 4)
            _safe_multi_cell(pdf, 5.5, f"-  {_sanitize_text(str(item))}")

    return bytes(pdf.output(dest="S"))