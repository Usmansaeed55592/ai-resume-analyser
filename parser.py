"""
Handles extracting raw text from an uploaded resume file (PDF or DOCX).
"""

import io
import pdfplumber
import docx


def extract_text_from_pdf(file) -> str:
    """Extract plain text from a PDF file-like object."""
    text_parts = []
    with pdfplumber.open(file) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
    return "\n".join(text_parts)


def extract_text_from_docx(file) -> str:
    """Extract plain text from a DOCX file-like object."""
    document = docx.Document(file)
    paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
    return "\n".join(paragraphs)


def extract_docx_paragraphs(file_bytes: bytes) -> list:
    """
    Extract non-empty paragraphs from a DOCX, keeping each paragraph's
    original index in the document. This index is later used to write
    the rewritten text back into the *same* paragraph object, so the
    original formatting (font, size, color, alignment) is preserved.

    Returns: list of {"index": int, "text": str}
    """
    document = docx.Document(io.BytesIO(file_bytes))
    paragraphs = []
    for idx, p in enumerate(document.paragraphs):
        text = p.text.strip()
        if text:
            paragraphs.append({"index": idx, "text": text})
    return paragraphs


def extract_resume_text(uploaded_file) -> str:
    """
    Dispatch to the correct extractor based on file extension.
    `uploaded_file` is expected to be a Streamlit UploadedFile object.
    """
    filename = uploaded_file.name.lower()

    if filename.endswith(".pdf"):
        return extract_text_from_pdf(uploaded_file)
    elif filename.endswith(".docx"):
        return extract_text_from_docx(uploaded_file)
    else:
        raise ValueError("Unsupported file type. Please upload a PDF or DOCX file.")


def clean_text(text: str) -> str:
    """Basic normalization: strip extra blank lines and whitespace."""
    lines = [line.strip() for line in text.splitlines()]
    lines = [line for line in lines if line]
    return "\n".join(lines)